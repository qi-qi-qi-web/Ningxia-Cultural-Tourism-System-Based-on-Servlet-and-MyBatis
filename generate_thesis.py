#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《文旅综合服务平台的设计与实现》毕业论文 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_before = Pt(0)
pf.space_after = Pt(0)

# 页面设置 A4
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_heading_custom(text, level=1, font_size=None, bold=True, center=False, space_before=0, space_after=0):
    """添加标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if font_size:
        run.font.size = Pt(font_size)
    else:
        sizes = {1: 16, 2: 14, 3: 13, 4: 12}
        run.font.size = Pt(sizes.get(level, 12))
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    pf = p.paragraph_format
    pf.space_before = Pt(space_before if space_before else (12 if level == 1 else 6))
    pf.space_after = Pt(space_after if space_after else 6)
    if center:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_body(text, indent=True, font_size=12, bold=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        pf.first_line_indent = Cm(0.74)  # 2字符缩进
    return p

def add_placeholder(text, font_size=12):
    """添加占位说明（用于截图位置）"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '楷体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    # 给占位符加上下划线边框效果
    return p

def add_table_from_data(headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置表头背景色
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


# ============================================================
# 封面
# ============================================================
# 空行
for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('文旅综合服务平台')
run.font.size = Pt(28)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('的设计与实现')
run.font.size = Pt(28)
run.bold = True
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('——宁夏文旅综合服务平台')
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(6):
    doc.add_paragraph()

# 封面底部信息
info_lines = [
    ('学    院：', '____________________'),
    ('专    业：', '____________________'),
    ('学生姓名：', '____________________'),
    ('学    号：', '____________________'),
    ('指导教师：', '____________________'),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}{value}')
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

# 分页
doc.add_page_break()

# ============================================================
# 中文摘要
# ============================================================
add_heading_custom('摘  要', level=1, font_size=18, center=True, space_before=24, space_after=12)

abstract_text = (
    '随着旅游业的蓬勃发展和人民生活水平的不断提高，人们对旅游体验和服务质量的需求日益增长，'
    '智慧旅游平台的建设得到了快速发展并引起了广泛关注。一个良好的文旅综合服务平台能方便管理员对数据进行快捷有效地管理，'
    '减少工作人员的大量重复性工作，提高效率、降低操作过程中的误差。'
)
add_body(abstract_text)

abstract_text2 = (
    '本系统采用 B/S 模式，使用 Java 编程语言，以 Servlet + JSP 结合 MyBatis 为技术架构，'
    '使用 MySQL 数据库。系统角色包括管理员和普通用户。管理员的主要功能包括景区管理、酒店管理、特产管理、'
    '攻略管理、用户管理、订单管理、新闻管理、公告管理、评论管理及日志管理等；普通用户的主要功能包括注册登录、'
    '浏览景点和酒店、购买特产、发布旅游攻略、发表评论以及收藏等。'
)
add_body(abstract_text2)

abstract_text3 = (
    '本文按照系统开发过程，依次介绍了该平台前期的准备工作，开发背景及所用的相关工具、技术，'
    '并对系统需求分析、系统设计以及系统实现等部分分别进行了阐述，最后对系统的完成情况做了总结与展望。'
)
add_body(abstract_text3)

p = doc.add_paragraph()
pf = p.paragraph_format
pf.space_before = Pt(12)
run = p.add_run('关键词：')
run.bold = True
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run = p.add_run('文旅服务；Java；MyBatis；JSP；MySQL')
run.font.size = Pt(12)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 英文摘要
doc.add_page_break()
add_heading_custom('Abstract', level=1, font_size=18, center=True, space_before=24, space_after=12)

eng_abstract = (
    'With the rapid development of the tourism industry and the continuous improvement of people\'s living standards, '
    'the demand for travel experience and service quality is growing. The construction of smart tourism platforms has '
    'developed rapidly and attracted widespread attention. A well-designed cultural tourism comprehensive service platform '
    'enables administrators to efficiently manage data, reduces repetitive tasks, enhances productivity, and minimizes operational errors.'
)
add_body(eng_abstract)

eng_abstract2 = (
    'This system adopts the B/S model, uses Java as the programming language, and employs Servlet + JSP combined with '
    'MyBatis as the technical framework, utilizing MySQL as the database. System roles include administrators and regular users. '
    'The main functions of administrators include scenic spot management, hotel management, specialty product management, '
    'travel guide management, user management, order management, news management, notice management, comment management, '
    'and log management. The main functions of regular users include registration and login, browsing scenic spots and hotels, '
    'purchasing specialty products, publishing travel guides, posting comments, and bookmarking favorites.'
)
add_body(eng_abstract2)

eng_abstract3 = (
    'Following the system development process, this paper introduces the preliminary preparation work, development background, '
    'and related tools and technologies used. It elaborates on system requirements analysis, system design, and system implementation, '
    'and finally provides a summary and outlook on the completion of the system.'
)
add_body(eng_abstract3)

p = doc.add_paragraph()
pf = p.paragraph_format
pf.space_before = Pt(12)
run = p.add_run('Key Words: ')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run = p.add_run('Cultural Tourism Service; Java; MyBatis; JSP; MySQL')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

# ============================================================
# 目录页（占位）
# ============================================================
doc.add_page_break()
add_heading_custom('目  录', level=1, font_size=18, center=True, space_before=24, space_after=12)
add_placeholder('【此处插入自动目录，在Word中使用"引用→目录"功能自动生成】')

# ============================================================
# 第1章 绪论
# ============================================================
doc.add_page_break()
add_heading_custom('第1章 绪论', level=1, font_size=18, center=True, space_before=24, space_after=12)

# 1.1
add_heading_custom('1.1 开发背景', level=2)

add_body(
    '随着我国经济的快速发展和人民生活水平的显著提高，旅游业已成为国民经济的重要组成部分。'
    '宁夏回族自治区拥有丰富的旅游资源，包括沙坡头、西夏王陵、沙湖、贺兰山岩画、镇北堡西部影城等知名景点，'
    '以及独特的回族文化风情和特色美食。然而，目前宁夏旅游信息化建设仍存在资源分散、信息不透明、服务不便捷等问题，'
    '游客难以一站式获取全面的旅游信息和服务。'
)
add_body(
    '近年来，"互联网+旅游"的发展模式得到广泛推广，智慧旅游平台的建设已成为提升旅游服务质量的重要手段。'
    '通过构建综合性的文旅服务平台，可以将景点、酒店、特产、攻略等旅游资源进行有效整合，为游客提供便捷的在线服务，'
    '同时为管理人员提供高效的数据管理工具。'
)
add_body(
    '在此背景下，本文设计并实现了一个面向宁夏地区的文旅综合服务平台，旨在通过信息化手段推动宁夏旅游产业的数字化升级。'
)

# 1.2
add_heading_custom('1.2 研究意义', level=2)

add_body(
    '本课题的研究意义主要体现在以下几个方面：'
)
add_body(
    '第一，整合宁夏旅游资源。系统将宁夏地区的景区、酒店、特产、旅游攻略等信息进行集中管理和展示，'
    '打破信息孤岛，为游客提供一站式的旅游信息服务。'
)
add_body(
    '第二，提升游客体验。游客可以通过平台方便地浏览景点信息、查看酒店详情、购买特产、发布和查看旅游攻略，'
    '并通过高德地图API实现路线规划与导航，极大提升了旅游出行的便利性。'
)
add_body(
    '第三，提高管理效率。管理员可以通过后台管理系统对平台数据进行统一管理，包括景点信息维护、订单处理、'
    '用户管理、评论审核等，有效减少了人工操作的重复性工作，降低了管理成本。'
)
add_body(
    '第四，促进宁夏旅游数字化。该平台的建设有助于推动宁夏旅游产业的信息化进程，为当地旅游经济发展提供技术支撑。'
)

# 1.3
add_heading_custom('1.3 国内外研究现状', level=2)

add_body(
    '在国外，旅游信息化建设起步较早，已形成了较为成熟的在线旅游服务体系。以TripAdvisor、Booking.com、'
    'Airbnb等为代表的国际旅游平台，提供了景点点评、酒店预订、民宿租赁等综合性服务。这些平台普遍采用先进的'
    '互联网技术，具有完善的功能体系和良好的用户体验。'
)
add_body(
    '在国内，随着互联网技术的快速发展，在线旅游平台也得到了蓬勃发展。携程旅行网提供了机票、酒店、'
    '旅游度假等全方位的旅行服务；马蜂窝旅游网以旅游攻略和社区互动为特色，吸引了大量用户分享旅游经验；'
    '美团等生活服务平台也逐步涉足旅游领域，提供了景点门票、酒店预订等服务。'
)
add_body(
    '然而，目前大多数旅游平台主要面向全国或全球市场，针对特定区域（如宁夏）的综合性文旅服务平台相对较少。'
    '宁夏拥有独特的自然景观和人文资源，但由于旅游信息化建设相对滞后，游客获取本地旅游信息的渠道较为有限。'
    '因此，开发一个面向宁夏的文旅综合服务平台具有重要的现实意义和应用价值。'
)

# 1.4
add_heading_custom('1.4 论文组织结构', level=2)

add_heading_custom('1.4.1 研究内容', level=3)
add_body(
    '本文围绕宁夏文旅综合服务平台的设计与实现展开研究，主要内容包括：系统需求分析、系统设计、系统实现与测试。'
    '在需求分析阶段，明确了系统的功能需求和非功能需求；在系统设计阶段，完成了系统架构设计、功能模块设计和数据库设计；'
    '在系统实现阶段，基于Java Servlet、JSP、MyBatis等核心技术完成了各功能模块的开发；最后对系统进行了功能测试。'
)

add_heading_custom('1.4.2 开发环境', level=3)
add_body('本系统的开发环境如下：')

headers_dev = ['项目', '配置']
rows_dev = [
    ['操作系统', 'Windows 10/11'],
    ['开发工具', 'IntelliJ IDEA'],
    ['编程语言', 'Java 25'],
    ['Web服务器', 'Apache Tomcat'],
    ['数据库', 'MySQL 5.7+ / MariaDB 10.2+'],
    ['后端框架', 'Java Servlet + MyBatis 3.5.19'],
    ['前端技术', 'JSP + Bootstrap 5 + jQuery'],
    ['地图API', '高德地图JavaScript API'],
    ['项目管理', 'Maven'],
    ['版本控制', 'Git'],
]
doc.add_paragraph()
add_table_from_data(headers_dev, rows_dev, col_widths=[4, 10])

# ============================================================
# 第2章 相关技术与工具
# ============================================================
doc.add_page_break()
add_heading_custom('第2章 相关技术与工具', level=1, center=True)

add_heading_custom('2.1 Java Web技术', level=2)
add_body(
    'Java是一种面向对象的编程语言，具有跨平台、安全性高、稳定性强等特点，'
    '广泛应用于企业级Web应用开发。Java Web开发基于Servlet和JSP技术，'
    '采用B/S（Browser/Server）架构模式，用户通过浏览器访问服务器上的Web应用，'
    '无需安装额外的客户端软件。'
)
add_body(
    'Servlet是Java Web应用的核心组件，负责处理客户端请求并生成响应。'
    'JSP（Java Server Pages）是一种动态网页技术，允许在HTML页面中嵌入Java代码，'
    '通过编译为Servlet来执行。本系统采用Servlet处理业务逻辑和请求分发，JSP负责前端页面展示。'
)

add_heading_custom('2.2 MyBatis框架', level=2)
add_body(
    'MyBatis是一款优秀的持久层框架，它支持自定义SQL、存储过程以及高级映射。'
    'MyBatis避免了几乎所有的JDBC代码和手动设置参数以及获取结果集的工作。'
    'MyBatis可以使用简单的XML或注解来配置和映射原生信息，将接口和Java的POJO（Plain Old Java Objects）'
    '映射成数据库中的记录。'
)
add_body(
    '本系统使用MyBatis作为数据访问层框架，通过XML映射文件定义SQL语句，'
    '实现了业务逻辑与数据访问的分离，提高了代码的可维护性和开发效率。'
    'MyBatis的灵活性和轻量级特点使其非常适合中小型Web应用的开发。'
)

add_heading_custom('2.3 MySQL数据库', level=2)
add_body(
    'MySQL是一个开源的关系型数据库管理系统，由瑞典MySQL AB公司开发，'
    '目前属于Oracle公司旗下产品。MySQL使用SQL（Structured Query Language）'
    '进行数据库管理，具有体积小、速度快、成本低、支持多种操作系统等优点，'
    '是中小型Web应用的首选数据库。'
)
add_body(
    '本系统采用MySQL 5.7+作为数据库管理系统，设计了15张数据库表，'
    '涵盖用户、景区、酒店、特产、攻略、评论、收藏、订单等核心业务数据。'
    '通过合理的表结构设计和索引优化，保证了系统数据的高效存取。'
)

add_heading_custom('2.4 JSP与Bootstrap', level=2)
add_body(
    'JSP（Java Server Pages）是一种基于Java的服务器端页面技术，它允许开发者在HTML中嵌入Java代码，'
    '通过JSP标签和表达式动态生成网页内容。JSP页面在第一次被访问时被编译为Servlet，'
    '后续请求直接执行编译后的代码，具有良好的性能表现。'
)
add_body(
    'Bootstrap是一个开源的前端框架，由Twitter公司开发，提供了丰富的CSS样式和JavaScript组件。'
    'Bootstrap基于响应式设计理念，能够自适应不同尺寸的屏幕设备，'
    '为系统提供了美观、一致的用户界面。本系统使用Bootstrap 5构建前端页面，'
    '结合自定义CSS样式，实现了良好的用户体验。'
)

add_heading_custom('2.5 高德地图API', level=2)
add_body(
    '高德地图开放平台提供了丰富的地图API服务，包括地图展示、定位、路线规划、导航等功能。'
    '本系统集成高德地图JavaScript API，在景点详情页面中嵌入地图组件，'
    '实现了景点位置的直观展示。同时，通过调用高德地图的路线规划接口，'
    '为游客提供了"到这里"的导航功能，支持公交、驾车、步行等多种出行方式的路线查询和价格参考，'
    '极大提升了游客的出行便利性。'
)

# ============================================================
# 第3章 系统需求分析
# ============================================================
doc.add_page_break()
add_heading_custom('第3章 系统需求分析', level=1, center=True)

add_heading_custom('3.1 可行性分析', level=2)

add_body(
    '技术可行性：本系统采用Java语言开发，基于B/S架构，使用成熟的Servlet、JSP、MyBatis和MySQL技术栈，'
    '这些技术已在实际项目中被广泛验证，具有成熟的技术生态和完善的社区支持。开发工具IntelliJ IDEA提供了强大的开发辅助功能，'
    '能够有效提高开发效率。因此，在技术上是可行的。'
)
add_body(
    '经济可行性：本系统所使用的开发工具和技术框架均为开源或免费版本，'
    '无需额外购买商业授权。系统部署所需的Web服务器（Tomcat）和数据库（MySQL）也是开源软件，'
    '运行成本较低。系统的维护和升级也相对简单，具有良好的经济可行性。'
)
add_body(
    '操作可行性：系统采用B/S架构，用户只需通过浏览器即可访问，无需安装额外软件。'
    '界面设计简洁清晰，操作流程符合用户习惯。管理员后台功能模块划分明确，'
    '管理人员经过简单培训即可熟练使用。因此，在操作上是可行的。'
)

add_heading_custom('3.2 功能需求分析', level=2)

add_heading_custom('3.2.1 系统角色分析', level=3)
add_body(
    '本系统的用户角色分为以下两类：'
)
add_body(
    '（1）普通用户：普通用户是指使用平台浏览和获取旅游服务的游客。普通用户可以注册和登录系统，'
    '浏览景点信息、酒店详情和特产商品，发布旅游攻略和评论，收藏感兴趣的内容，'
    '下单购买特产，以及管理个人资料。'
)
add_body(
    '（2）管理员：管理员负责平台后台的维护和管理工作。管理员可以管理景点信息、酒店信息、特产信息、'
    '旅游攻略、用户账户、处理订单、审核评论、管理新闻公告以及查看系统日志等。'
)

add_heading_custom('3.2.2 系统业务流程分析', level=3)
add_placeholder('【此处插入系统业务流程图】')
add_body(
    '系统业务流程主要包括用户操作流程和管理员操作流程。用户操作流程为：用户注册登录后，'
    '可以在首页浏览推荐景点，查看景点详情并使用地图导航功能，浏览酒店信息，'
    '在特产商城选购商品并下单，在攻略中心发布和浏览旅游攻略，对感兴趣的内容进行评论和收藏。'
    '管理员操作流程为：管理员登录后台管理系统，通过各功能模块对平台数据进行增删改查操作。'
)

add_heading_custom('3.2.3 系统用例分析', level=3)
add_placeholder('【此处插入系统用例图】')

add_body('系统的核心用例包括：', indent=False)

add_body('（1）用户注册与登录：用户通过注册功能创建账户，注册后使用用户名和密码登录系统。')
add_body('（2）景点浏览与导航：用户浏览景点列表和详情，查看景点位置地图，使用"到这里"功能进行路线规划。')
add_body('（3）酒店浏览：用户查看酒店列表和详情，了解房型、价格和设施信息。')
add_body('（4）特产购买：用户浏览特产商品，加入购物车，提交订单并查看订单状态。')
add_body('（5）攻略发布：用户编写并发布旅游攻略，可添加标签和图片。')
add_body('（6）评论与收藏：用户对景点、酒店、特产、攻略等内容进行评论或收藏。')
add_body('（7）景点管理：管理员对景区信息进行增删改查操作，包括经纬度、门票价格等。')
add_body('（8）酒店管理：管理员对酒店和客房信息进行管理。')
add_body('（9）特产管理：管理员对特产分类和商品信息进行管理。')
add_body('（10）用户管理：管理员对用户账户进行管理，包括添加、编辑、禁用等操作。')
add_body('（11）订单管理：管理员查看和处理用户订单，更新订单状态。')
add_body('（12）攻略管理：管理员审核和管理用户发布的旅游攻略。')
add_body('（13）评论管理：管理员审核和管理用户的评论内容。')
add_body('（14）日志管理：管理员查看系统操作日志，了解平台运行情况。')

add_heading_custom('3.3 非功能需求分析', level=2)
add_body(
    '（1）系统性能：系统应能够支持一定数量的并发用户访问，页面响应时间在合理范围内（一般不超过3秒），'
    '数据库查询效率应满足日常使用需求。'
)
add_body(
    '（2）安全性：用户密码应经过加密存储（本系统采用SHA-256算法），防止密码泄露。'
    '系统应具备用户会话管理机制，防止未授权访问。不同角色的用户应具有不同的操作权限。'
)
add_body(
    '（3）可扩展性：系统架构应具有良好的可扩展性，便于后续增加新的功能模块。'
    '数据库设计应考虑到未来业务扩展的需要，预留适当的扩展空间。'
)
add_body(
    '（4）可用性：系统界面应简洁美观、操作直观，符合用户的使用习惯。'
    '系统应提供必要的错误提示和操作引导，降低用户的学习成本。'
)

# ============================================================
# 第4章 系统设计
# ============================================================
doc.add_page_break()
add_heading_custom('第4章 系统设计', level=1, center=True)

add_heading_custom('4.1 总体设计', level=2)

add_heading_custom('4.1.1 系统架构设计', level=3)
add_placeholder('【此处插入系统架构图】')
add_body(
    '本系统采用B/S（Browser/Server）三层架构，即表现层、业务逻辑层和数据访问层。'
    '表现层负责与用户交互，使用JSP和Bootstrap构建前端页面，通过浏览器向服务器发送请求并展示响应结果。'
    '业务逻辑层由Java Servlet组成，负责处理业务逻辑，接收表现层请求，调用数据访问层接口获取数据，'
    '并将结果返回给表现层。数据访问层使用MyBatis框架与MySQL数据库进行交互，'
    '通过XML映射文件定义SQL语句，实现数据的增删改查操作。'
)
add_body(
    '这种三层架构的设计模式将界面展示、业务处理和数据访问分离，'
    '降低了各层之间的耦合度，提高了系统的可维护性和可扩展性。'
)

add_heading_custom('4.1.2 功能模块设计', level=3)
add_placeholder('【此处插入系统功能模块图】')

add_body('本系统功能模块分为前台模块和后台管理模块两大部分：', indent=False)

add_body('前台功能模块包括：')
add_body('（1）首页展示模块：展示热门景点推荐、系统公告、旅游攻略精选等内容。')
add_body('（2）景点浏览模块：景点列表展示、景点详情查看、地图导航、门票信息。')
add_body('（3）酒店浏览模块：酒店列表展示、酒店详情查看、房型信息及设施。')
add_body('（4）特产商城模块：特产分类浏览、商品详情、购物车、订单管理。')
add_body('（5）攻略中心模块：攻略列表、攻略详情、攻略发布与管理。')
add_body('（6）个人中心模块：个人信息管理、我的收藏、我的评论、我的订单。')
add_body('（7）注册登录模块：用户注册、登录、找回密码。')

add_body('后台管理模块包括：')
add_body('（1）景区管理：景点信息增删改查，图片管理。')
add_body('（2）酒店管理：酒店及客房信息管理。')
add_body('（3）特产管理：特产分类和商品信息管理。')
add_body('（4）用户管理：用户账户管理和权限设置。')
add_body('（5）订单管理：订单状态管理和处理。')
add_body('（6）攻略管理：旅游攻略审核和标签管理。')
add_body('（7）评论管理：用户评论审核和管理。')
add_body('（8）新闻管理：新闻动态发布和管理。')
add_body('（9）公告管理：通知公告发布和管理。')
add_body('（10）日志管理：系统操作日志查看。')

add_heading_custom('4.1.3 系统安全设计', level=3)
add_body(
    '在系统安全方面，本系统采取了以下措施：'
)
add_body(
    '（1）密码加密：用户密码通过SHA-256算法进行哈希处理后存储，即使数据库泄露，'
    '攻击者也难以还原原始密码。在登录验证时，将用户输入的密码经过相同的哈希处理，'
    '与数据库中的哈希值进行比对。'
)
add_body(
    '（2）会话管理：用户登录后，系统创建会话（Session）保存用户登录状态，'
    '并通过Cookie机制维持会话。在访问需要权限的页面时，系统会检查用户是否已登录，'
    '防止未授权访问。'
)
add_body(
    '（3）角色权限控制：系统通过用户角色字段区分普通用户和管理员。'
    '对于管理员后台操作，系统会验证当前用户是否具有管理员权限，'
    '确保只有授权管理员才能执行敏感操作。'
)

add_heading_custom('4.2 数据库设计', level=2)

add_heading_custom('4.2.1 数据库E-R图设计', level=3)
add_placeholder('【此处插入数据库E-R图】')
add_body(
    '根据系统功能需求，数据库主要包含以下实体：用户（User）、景点（ScenicSpot）、'
    '门票（ScenicTicket）、酒店（Hotel）、客房（HotelRoom）、特产（Specialty）、'
    '特产分类（SpecialtyCategory）、攻略（TravelGuide）、攻略标签（GuideTag）、'
    '评论（Comment）、收藏（Favorite）、订单主表（OrderMain）、订单明细（OrderItem）、'
    '新闻动态（NewsDynamic）、通知公告（OfficialNotice）和平台日志（PlatformLog），'
    '共计16张数据表。'
)
add_body(
    '各实体之间的关系如下：一个景点对应多个门票类型；一个酒店对应多个客房；'
    '一个特产分类对应多个特产商品；一篇攻略可以有多个标签，标签既可以作为模板独立存在'
    '（guide_id为NULL），也可以关联到具体攻略；一个用户可以对多个内容发表评论或收藏；'
    '一个订单包含多个订单明细项。'
)

add_heading_custom('4.2.2 数据表设计', level=3)
add_body(
    '根据E-R图设计，本系统在MySQL数据库中创建了以下16张数据表，'
    '涵盖用户、景区、特产、酒店、攻略、标签、订单、评论、收藏、日志、新闻动态、通知公告等核心业务数据。'
    '以下列出所有数据表的详细结构，包括字段名、数据类型、说明及完整性约束。'
)

# ============================================================
# 辅助函数：生成带完整性约束的表
# ============================================================
def add_table_with_constraints(table_label, table_name, table_comment, field_data):
    """
    参数:
      table_label: 表编号，如 '4-1'
      table_name: 表名，如 'sys_user'
      table_comment: 中文注释
      field_data: list of (字段名, 类型, 说明, 完整性约束)
    """
    add_body(f'表{table_label} {table_name}（{table_comment}）', indent=False, bold=True)
    headers = ['字段名', '数据类型', '说明', '完整性约束']
    rows = [[f[0], f[1], f[2], f[3]] for f in field_data]
    add_table_from_data(headers, rows, col_widths=[2.8, 3.0, 4.5, 4.2])

# ========================================================================
# 第一部分：用户层 (2张表)
# ========================================================================

# 1. sys_user
add_table_with_constraints('4-1', 'sys_user', '用户表', [
    ('id',              'BIGINT',           '用户ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('username',        'VARCHAR(50)',      '用户名',               'NOT NULL, UNIQUE'),
    ('password',        'VARCHAR(255)',     '密码（SHA-256加密存储）','NOT NULL'),
    ('nickname',        'VARCHAR(50)',      '昵称',                 'DEFAULT NULL'),
    ('email',           'VARCHAR(100)',     '邮箱',                 'DEFAULT NULL'),
    ('phone',           'VARCHAR(20)',      '手机号',               'DEFAULT NULL'),
    ('avatar',          'VARCHAR(500)',     '头像URL',              'DEFAULT NULL'),
    ('role',            "ENUM('USER','ADMIN')", '角色: USER=普通用户, ADMIN=管理员', "NOT NULL, DEFAULT 'USER'"),
    ('status',          'TINYINT',          '状态: 0=禁用, 1=正常', "NOT NULL, DEFAULT 1"),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
    ('updated_at',      'DATETIME',         '更新时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP, ON UPDATE CURRENT_TIMESTAMP"),
])

# 2. sys_platform_log
add_table_with_constraints('4-2', 'sys_platform_log', '平台日志表（用户活动记录）', [
    ('id',              'BIGINT',           '日志ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('user_id',         'BIGINT',           '用户ID',               'NOT NULL, FOREIGN KEY REFERENCES sys_user(id) ON DELETE CASCADE'),
    ('user_name',       'VARCHAR(50)',      '用户名（冗余）',        'NOT NULL'),
    ('log_type',        "ENUM('REGISTER','LOGIN','LOGOUT','POST_COMMENT','POST_GUIDE','PLACE_ORDER','CANCEL_ORDER','CONFIRM_RECEIPT','UPDATE_PROFILE')",
                                                                     '日志类型', "NOT NULL"),
    ('target_type',     'VARCHAR(50)',      '关联目标类型',          "DEFAULT NULL"),
    ('target_id',       'BIGINT',           '关联目标记录ID',        'DEFAULT NULL'),
    ('target_name',     'VARCHAR(200)',     '关联目标名称（冗余）',   'DEFAULT NULL'),
    ('detail',          'JSON',             '详细信息: 如IP、备注等',  'DEFAULT NULL'),
    ('ip_address',      'VARCHAR(50)',      '操作来源IP地址',        'DEFAULT NULL'),
    ('user_agent',      'VARCHAR(500)',     '浏览器User-Agent',     'DEFAULT NULL'),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
])

# ========================================================================
# 第二部分：核心业务层 (9张表)
# ========================================================================

# 3. scenic_spot
add_table_with_constraints('4-3', 'scenic_spot', '景区表', [
    ('id',              'BIGINT',           '景区ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('name',            'VARCHAR(100)',     '景区名称',             'NOT NULL'),
    ('description',     'TEXT',             '景区介绍',             'DEFAULT NULL'),
    ('address',         'VARCHAR(255)',     '详细地址',             'DEFAULT NULL'),
    ('province',        'VARCHAR(50)',      '省',                   'DEFAULT NULL'),
    ('city',            'VARCHAR(50)',      '市',                   'DEFAULT NULL'),
    ('district',        'VARCHAR(50)',      '区/县',                'DEFAULT NULL'),
    ('latitude',        'DECIMAL(10,7)',    '纬度',                 'DEFAULT NULL'),
    ('longitude',       'DECIMAL(10,7)',    '经度',                 'DEFAULT NULL'),
    ('opening_hours',   'VARCHAR(200)',     '开放时间描述',         'DEFAULT NULL'),
    ('contact_phone',   'VARCHAR(20)',      '景区联系电话',         'DEFAULT NULL'),
    ('cover_image',     'VARCHAR(500)',     '封面图URL',            'DEFAULT NULL'),
    ('images',          'JSON',             '图片列表JSON',         'DEFAULT NULL'),
    ('min_price',       'DECIMAL(10,2)',    '最低门票价',           'DEFAULT NULL'),
    ('view_count',      'BIGINT',           '浏览量',               "NOT NULL, DEFAULT 0"),
    ('favorite_count',  'BIGINT',           '收藏数（冗余计数）',    "NOT NULL, DEFAULT 0"),
    ('status',          "ENUM('OPEN','CLOSED','MAINTENANCE')", '运营状态', "NOT NULL, DEFAULT 'OPEN'"),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
    ('updated_at',      'DATETIME',         '更新时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP, ON UPDATE CURRENT_TIMESTAMP"),
])

# 4. scenic_ticket
add_table_with_constraints('4-4', 'scenic_ticket', '景区门票表', [
    ('id',              'BIGINT',           '门票ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('scenic_spot_id',  'BIGINT',           '所属景区ID',           'NOT NULL, FOREIGN KEY REFERENCES scenic_spot(id) ON DELETE CASCADE'),
    ('ticket_name',     'VARCHAR(100)',     '票种名称',             'NOT NULL'),
    ('ticket_type',     "ENUM('ADULT','CHILD','STUDENT','SENIOR','GROUP')",
                                                                     '票种类型', "NOT NULL"),
    ('price',           'DECIMAL(10,2)',    '原价',                 'NOT NULL'),
    ('discount_price',  'DECIMAL(10,2)',    '折后价（活动价）',     'DEFAULT NULL'),
    ('stock',           'INT',              '每日库存限额',          "NOT NULL, DEFAULT 0"),
    ('description',     'TEXT',             '购买须知/票种说明',    'DEFAULT NULL'),
    ('status',          'TINYINT',          '状态: 1=上架, 0=下架', "NOT NULL, DEFAULT 1"),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
    ('updated_at',      'DATETIME',         '更新时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP, ON UPDATE CURRENT_TIMESTAMP"),
])

# 5. news_dynamic
add_table_with_constraints('4-5', 'news_dynamic', '新闻动态表（民间视角）', [
    ('id',              'BIGINT',           '新闻ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('title',           'VARCHAR(200)',     '标题',                 'NOT NULL'),
    ('content',         'LONGTEXT',         '正文内容（富文本）',    'NOT NULL'),
    ('cover_image',     'VARCHAR(500)',     '封面图',               'DEFAULT NULL'),
    ('source',          'VARCHAR(100)',     '来源',                 'DEFAULT NULL'),
    ('author_name',     'VARCHAR(50)',      '投稿人/作者',          'DEFAULT NULL'),
    ('view_count',      'INT',              '浏览量',               "NOT NULL, DEFAULT 0"),
    ('is_published',    'TINYINT',          '是否已发布: 1=是, 0=草稿', "NOT NULL, DEFAULT 0"),
    ('published_at',    'DATETIME',         '发布时间',             'DEFAULT NULL'),
    ('created_by',      'BIGINT',           '发布人（管理员ID）',   'DEFAULT NULL, FOREIGN KEY REFERENCES sys_user(id) ON DELETE SET NULL'),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
    ('updated_at',      'DATETIME',         '更新时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP, ON UPDATE CURRENT_TIMESTAMP"),
])

# 6. official_notice
add_table_with_constraints('4-6', 'official_notice', '通知公告表（官方视角）', [
    ('id',              'BIGINT',           '公告ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('title',           'VARCHAR(200)',     '标题',                 'NOT NULL'),
    ('content',         'LONGTEXT',         '正文内容（富文本）',    'NOT NULL'),
    ('scenic_spot_id',  'BIGINT',           '关联景区ID',           'DEFAULT NULL, FOREIGN KEY REFERENCES scenic_spot(id) ON DELETE SET NULL'),
    ('cover_image',     'VARCHAR(500)',     '封面图',               'DEFAULT NULL'),
    ('is_top',          'TINYINT',          '是否置顶: 1=是, 0=否', "NOT NULL, DEFAULT 0"),
    ('is_published',    'TINYINT',          '是否已发布: 1=是, 0=草稿', "NOT NULL, DEFAULT 0"),
    ('published_at',    'DATETIME',         '发布时间',             'DEFAULT NULL'),
    ('created_by',      'BIGINT',           '发布人（管理员ID）',   'DEFAULT NULL, FOREIGN KEY REFERENCES sys_user(id) ON DELETE SET NULL'),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
    ('updated_at',      'DATETIME',         '更新时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP, ON UPDATE CURRENT_TIMESTAMP"),
])

# 7. specialty_category
add_table_with_constraints('4-7', 'specialty_category', '特产分类表', [
    ('id',              'BIGINT',           '分类ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('name',            'VARCHAR(50)',      '分类名称',             'NOT NULL'),
    ('description',     'VARCHAR(255)',     '分类描述',             'DEFAULT NULL'),
    ('icon',            'VARCHAR(500)',     '分类图标URL',          'DEFAULT NULL'),
    ('sort_order',      'INT',              '排序（越小越靠前）',    "NOT NULL, DEFAULT 0"),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
])

# 8. specialty
add_table_with_constraints('4-8', 'specialty', '本地特产表', [
    ('id',              'BIGINT',           '特产ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('category_id',     'BIGINT',           '所属分类ID',           'NOT NULL, FOREIGN KEY REFERENCES specialty_category(id) ON DELETE RESTRICT'),
    ('scenic_spot_id',  'BIGINT',           '产地景区ID',           'DEFAULT NULL, FOREIGN KEY REFERENCES scenic_spot(id) ON DELETE SET NULL'),
    ('name',            'VARCHAR(100)',     '特产名称',             'NOT NULL'),
    ('description',     'TEXT',             '特产描述',             'DEFAULT NULL'),
    ('price',           'DECIMAL(10,2)',    '价格',                 'NOT NULL'),
    ('stock',           'INT',              '库存',                 "NOT NULL, DEFAULT 0"),
    ('main_image',      'VARCHAR(500)',     '主图URL',              'DEFAULT NULL'),
    ('images',          'JSON',             '图片列表JSON',         'DEFAULT NULL'),
    ('sales_count',     'INT',              '累计销量',             "NOT NULL, DEFAULT 0"),
    ('favorite_count',  'BIGINT',           '收藏数（冗余计数）',    "NOT NULL, DEFAULT 0"),
    ('status',          'TINYINT',          '状态: 1=上架, 0=下架', "NOT NULL, DEFAULT 1"),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
    ('updated_at',      'DATETIME',         '更新时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP, ON UPDATE CURRENT_TIMESTAMP"),
])

# 9. hotel
add_table_with_constraints('4-9', 'hotel', '民宿酒店表', [
    ('id',              'BIGINT',           '酒店ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('name',            'VARCHAR(100)',     '酒店名称',             'NOT NULL'),
    ('description',     'TEXT',             '酒店介绍',             'DEFAULT NULL'),
    ('address',         'VARCHAR(255)',     '详细地址',             'DEFAULT NULL'),
    ('province',        'VARCHAR(50)',      '省',                   'DEFAULT NULL'),
    ('city',            'VARCHAR(50)',      '市',                   'DEFAULT NULL'),
    ('district',        'VARCHAR(50)',      '区/县',                'DEFAULT NULL'),
    ('latitude',        'DECIMAL(10,7)',    '纬度',                 'DEFAULT NULL'),
    ('longitude',       'DECIMAL(10,7)',    '经度',                 'DEFAULT NULL'),
    ('star_rating',     'TINYINT',          '星级(1-5)',            'DEFAULT NULL'),
    ('contact_phone',   'VARCHAR(20)',      '联系电话',             'DEFAULT NULL'),
    ('facilities',      'JSON',             '设施列表JSON',         'DEFAULT NULL'),
    ('cover_image',     'VARCHAR(500)',     '封面图',               'DEFAULT NULL'),
    ('images',          'JSON',             '图片列表JSON',         'DEFAULT NULL'),
    ('min_price',       'DECIMAL(10,2)',    '最低房价（列表展示用）','DEFAULT NULL'),
    ('favorite_count',  'BIGINT',           '收藏数（冗余计数）',    "NOT NULL, DEFAULT 0"),
    ('status',          'TINYINT',          '状态: 1=营业, 0=歇业', "NOT NULL, DEFAULT 1"),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
    ('updated_at',      'DATETIME',         '更新时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP, ON UPDATE CURRENT_TIMESTAMP"),
])

# 10. hotel_room
add_table_with_constraints('4-10', 'hotel_room', '酒店房间表', [
    ('id',              'BIGINT',           '房间ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('hotel_id',        'BIGINT',           '所属酒店ID',           'NOT NULL, FOREIGN KEY REFERENCES hotel(id) ON DELETE CASCADE'),
    ('room_name',       'VARCHAR(100)',     '房间名称',             'NOT NULL'),
    ('room_type',       'VARCHAR(50)',      '房型: 大床房/双床房/套房/家庭房', 'NOT NULL'),
    ('price',           'DECIMAL(10,2)',    '原价（每晚）',         'NOT NULL'),
    ('discount_price',  'DECIMAL(10,2)',    '折后价（每晚）',       'DEFAULT NULL'),
    ('stock',           'INT',              '可售房间数',           "NOT NULL, DEFAULT 0"),
    ('max_guests',      'TINYINT',          '最大入住人数',         "NOT NULL, DEFAULT 2"),
    ('bed_type',        'VARCHAR(50)',      '床型描述',             'DEFAULT NULL'),
    ('area',            'DECIMAL(6,2)',     '房间面积(m²)',         'DEFAULT NULL'),
    ('facilities',      'JSON',             '房间设施JSON',         'DEFAULT NULL'),
    ('images',          'JSON',             '房间图片JSON',         'DEFAULT NULL'),
    ('status',          'TINYINT',          '状态: 1=可订, 0=不可订', "NOT NULL, DEFAULT 1"),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
    ('updated_at',      'DATETIME',         '更新时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP, ON UPDATE CURRENT_TIMESTAMP"),
])

# 11. travel_guide
add_table_with_constraints('4-11', 'travel_guide', '旅游攻略表', [
    ('id',              'BIGINT',           '攻略ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('user_id',         'BIGINT',           '发布者ID',             'NOT NULL, FOREIGN KEY REFERENCES sys_user(id) ON DELETE CASCADE'),
    ('title',           'VARCHAR(200)',     '攻略标题',             'NOT NULL'),
    ('content',         'LONGTEXT',         '攻略正文（富文本）',    'NOT NULL'),
    ('cover_image',     'VARCHAR(500)',     '封面图',               'DEFAULT NULL'),
    ('tags',            'VARCHAR(500)',     '标签（逗号分隔）',     'DEFAULT NULL'),
    ('like_count',      'INT',              '点赞数',               "NOT NULL, DEFAULT 0"),
    ('view_count',      'INT',              '浏览数',               "NOT NULL, DEFAULT 0"),
    ('comment_count',   'INT',              '评论数（冗余计数）',    "NOT NULL, DEFAULT 0"),
    ('favorite_count',  'BIGINT',           '收藏数（冗余计数）',    "NOT NULL, DEFAULT 0"),
    ('status',          "ENUM('PUBLISHED','DRAFT','HIDDEN')", '状态', "NOT NULL, DEFAULT 'PUBLISHED'"),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
    ('updated_at',      'DATETIME',         '更新时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP, ON UPDATE CURRENT_TIMESTAMP"),
])

# 12. guide_tag
add_table_with_constraints('4-12', 'guide_tag', '攻略标签表（含攻略关联）', [
    ('id',              'BIGINT',           '标签ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('guide_id',        'BIGINT',           '关联攻略ID（NULL=模板标签）', 'DEFAULT NULL, FOREIGN KEY REFERENCES travel_guide(id) ON DELETE CASCADE'),
    ('name',            'VARCHAR(50)',      '标签名称',             'NOT NULL'),
    ('category',        "ENUM('FEATURE','TIME','AUDIENCE','BUDGET')",
                                                                     '标签分类: FEATURE=特点, TIME=时间, AUDIENCE=适合人群, BUDGET=预算', "NOT NULL"),
    ('sort_order',      'INT',              '排序序号',             'DEFAULT 0'),
])

# ========================================================================
# 第三部分：交互交易层 (5张表)
# ========================================================================

# 13. order_main
add_table_with_constraints('4-13', 'order_main', '订单主表（特产订单）', [
    ('id',              'BIGINT',           '订单ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('order_no',        'VARCHAR(32)',      '订单编号',             'NOT NULL, UNIQUE'),
    ('user_id',         'BIGINT',           '下单用户ID',           'NOT NULL, FOREIGN KEY REFERENCES sys_user(id) ON DELETE RESTRICT'),
    ('total_amount',    'DECIMAL(10,2)',    '订单总额',             'NOT NULL'),
    ('discount_amount', 'DECIMAL(10,2)',    '优惠金额',             "NOT NULL, DEFAULT 0.00"),
    ('pay_amount',      'DECIMAL(10,2)',    '实付金额',             'NOT NULL'),
    ('pickup_method',   "ENUM('PICKUP','DELIVERY')", '取货方式', "NOT NULL, DEFAULT 'PICKUP'"),
    ('delivery_name',   'VARCHAR(50)',      '收件人姓名',           'DEFAULT NULL'),
    ('delivery_phone',  'VARCHAR(20)',      '收件人电话',           'DEFAULT NULL'),
    ('delivery_address','VARCHAR(255)',     '收件地址',             'DEFAULT NULL'),
    ('status',          "ENUM('PENDING','PAID','SHIPPED','CANCELLED','REFUNDED','COMPLETED')",
                                                                     '订单状态', "NOT NULL, DEFAULT 'PENDING'"),
    ('payment_method',  'VARCHAR(50)',      '支付方式',             'DEFAULT NULL'),
    ('paid_at',         'DATETIME',         '支付时间',             'DEFAULT NULL'),
    ('shipped_at',      'DATETIME',         '发货时间',             'DEFAULT NULL'),
    ('completed_at',    'DATETIME',         '确认收货时间',         'DEFAULT NULL'),
    ('remark',          'VARCHAR(500)',     '用户备注',             'DEFAULT NULL'),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
    ('updated_at',      'DATETIME',         '更新时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP, ON UPDATE CURRENT_TIMESTAMP"),
])

# 13. order_item
add_table_with_constraints('4-14', 'order_item', '订单明细表（特产）', [
    ('id',              'BIGINT',           '明细ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('order_id',        'BIGINT',           '所属订单ID',           'NOT NULL, FOREIGN KEY REFERENCES order_main(id) ON DELETE CASCADE'),
    ('specialty_id',    'BIGINT',           '特产ID',               'NOT NULL, FOREIGN KEY REFERENCES specialty(id) ON DELETE RESTRICT'),
    ('item_name',       'VARCHAR(200)',     '商品名称（冗余）',     'NOT NULL'),
    ('item_image',      'VARCHAR(500)',     '商品图片（冗余）',     'DEFAULT NULL'),
    ('quantity',        'INT',              '购买数量',             "NOT NULL, DEFAULT 1"),
    ('unit_price',      'DECIMAL(10,2)',    '单价',                 'NOT NULL'),
    ('subtotal',        'DECIMAL(10,2)',    '小计',                 'NOT NULL'),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
])

# 14. comment
add_table_with_constraints('4-15', 'comment', '评论表（多态关联：景区/特产/酒店/攻略）', [
    ('id',              'BIGINT',           '评论ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('user_id',         'BIGINT',           '评论用户ID',           'NOT NULL, FOREIGN KEY REFERENCES sys_user(id) ON DELETE CASCADE'),
    ('target_type',     "ENUM('SCENIC','SPECIALTY','HOTEL','GUIDE')",
                                                                     '评论目标类型', "NOT NULL"),
    ('target_id',       'BIGINT',           '评论目标ID',           'NOT NULL'),
    ('content',         'TEXT',             '评论内容',             'NOT NULL'),
    ('rating',          'TINYINT',          '评分(1-5)',            'DEFAULT NULL'),
    ('images',          'JSON',             '评论图片JSON',         'DEFAULT NULL'),
    ('parent_id',       'BIGINT',           '父评论ID（支持楼中楼回复）', 'DEFAULT NULL, FOREIGN KEY REFERENCES comment(id) ON DELETE CASCADE'),
    ('like_count',      'INT',              '点赞数',               "NOT NULL, DEFAULT 0"),
    ('status',          'TINYINT',          '状态: 1=显示, 0=隐藏', "NOT NULL, DEFAULT 1"),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
    ('updated_at',      'DATETIME',         '更新时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP, ON UPDATE CURRENT_TIMESTAMP"),
])

# 15. favorite
add_table_with_constraints('4-16', 'favorite', '收藏表（多态关联：景区/特产/酒店/攻略）', [
    ('id',              'BIGINT',           '收藏ID',               'PRIMARY KEY, AUTO_INCREMENT'),
    ('user_id',         'BIGINT',           '用户ID',               'NOT NULL, FOREIGN KEY REFERENCES sys_user(id) ON DELETE CASCADE'),
    ('target_type',     "ENUM('SCENIC','SPECIALTY','HOTEL','GUIDE')",
                                                                     '收藏目标类型', "NOT NULL"),
    ('target_id',       'BIGINT',           '收藏目标ID',           'NOT NULL'),
    ('created_at',      'DATETIME',         '创建时间',             "NOT NULL, DEFAULT CURRENT_TIMESTAMP"),
])

# ============================================================
# 第5章 系统实现
# ============================================================
doc.add_page_break()
add_heading_custom('第5章 系统实现', level=1, center=True)

add_heading_custom('5.1 登录注册模块实现', level=2)

add_heading_custom('5.1.1 用户注册', level=3)
add_body(
    '用户注册是系统的基础功能。用户在注册页面填写用户名、密码、邮箱、手机号等信息，'
    '提交后系统会对用户名进行唯一性校验，并对密码进行SHA-256加密后存入数据库。'
    '注册成功后，用户可以使用注册的账号和密码登录系统。'
)
add_placeholder('【此处插入用户注册页面截图】')

add_heading_custom('5.1.2 用户登录', level=3)
add_body(
    '用户登录功能验证用户输入的账号和密码是否正确。系统将用户输入的密码进行SHA-256哈希加密，'
    '与数据库中存储的哈希值进行比对。验证通过后，系统创建会话（Session）记录用户登录状态，'
    '并跳转至系统首页。登录功能同时支持普通用户和管理员。'
)
add_placeholder('【此处插入用户登录页面截图】')

add_heading_custom('5.1.3 密码加密实现', level=3)
add_body(
    '本系统使用SHA-256算法对用户密码进行加密。在用户注册时，系统将明文密码通过SHA-256哈希算法'
    '转换为固定长度的哈希值，并在哈希值前添加"$SHA$"前缀标识加密方式，然后存入数据库。'
    '在用户登录时，系统对输入密码进行同样的哈希处理，与数据库存储的哈希值进行比对，'
    '从而验证密码的正确性。这种方式确保了即使数据库泄露，攻击者也难以还原出用户的原始密码。'
)

add_heading_custom('5.2 前台功能实现', level=2)

add_heading_custom('5.2.1 首页展示', level=3)
add_body(
    '系统首页是用户访问平台的第一界面。首页顶部为导航栏，包含景点、酒店、特产、攻略等主要功能入口。'
    '首页主体区域展示热门景点轮播图、推荐景点列表、最新旅游攻略以及系统公告等内容。'
    '页面设计采用Bootstrap框架，布局清晰，视觉美观。'
)
add_placeholder('【此处插入系统首页截图】')

add_heading_custom('5.2.2 景点浏览与详情', level=3)
add_body(
    '景点浏览模块展示宁夏各景点的列表信息。用户点击景点卡片可进入详情页面，'
    '查看景点的详细介绍、开放时间、地址、门票价格等信息。详情页面还集成了高德地图组件，'
    '直观展示景点的地理位置。用户点击"到这里"按钮，可以调用高德地图API进行路线规划，'
    '获取公交、驾车、步行等出行方式的路线信息和价格参考。'
)
add_placeholder('【此处插入景点详情页面截图】')
add_placeholder('【此处插入地图导航功能截图】')

add_heading_custom('5.2.3 酒店浏览', level=3)
add_body(
    '酒店模块展示宁夏地区的酒店和民宿信息。酒店列表以卡片形式展示酒店名称、星级、地址和参考价格。'
    '用户点击进入酒店详情页，可以查看酒店的详细介绍、设施服务以及各房型的价格信息，'
    '方便用户根据需求选择合适的住宿。'
)
add_placeholder('【此处插入酒店浏览页面截图】')

add_heading_custom('5.2.4 特产商城与订单', level=3)
add_body(
    '特产商城模块展示宁夏当地的特色商品，如枸杞、八宝茶、贺兰石等。商品按类别分类展示，'
    '用户可以将心仪的商品加入购物车，填写收货信息后提交订单。系统支持快递配送和到店自取两种方式。'
    '用户可以在个人中心查看订单状态，包括待付款、已付款、已发货、已完成和已取消等状态。'
)
add_placeholder('【此处插入特产商城页面截图】')
add_placeholder('【此处插入订单管理页面截图】')

add_heading_custom('5.2.5 攻略发布与浏览', level=3)
add_body(
    '攻略中心是用户分享旅游经验和心得的功能模块。用户可以浏览其他用户发布的旅游攻略，'
    '也可以自己撰写和发布攻略。发布攻略时，用户可以添加标题、正文内容、封面图片和标签。'
    '攻略支持点赞、评论和收藏功能，增加了用户之间的互动性。'
)
add_placeholder('【此处插入攻略中心页面截图】')

add_heading_custom('5.2.6 评论与收藏', level=3)
add_body(
    '评论功能允许用户对景点、酒店、特产和攻略等内容发表看法和评价。评论采用多态设计，'
    '通过target_type字段区分不同的评论目标类型。评论支持回复功能，用户之间可以进行互动交流。'
    '收藏功能允许用户将感兴趣的内容保存到个人中心，方便日后快速查看。'
)
add_placeholder('【此处插入评论功能截图】')

add_heading_custom('5.3 后台管理模块实现', level=2)

add_heading_custom('5.3.1 景点管理', level=3)
add_body(
    '景点管理模块是管理员维护景区信息的功能界面。管理员可以添加新的景点，编辑已存在的景点信息，'
    '包括景点名称、描述、地址、经纬度、开放时间、门票价格和图片等。'
    '经纬度字段为必填项，确保地图导航功能的正常使用。管理员还可以管理景点相册图片。'
)
add_placeholder('【此处插入景点管理页面截图】')

add_heading_custom('5.3.2 酒店管理', level=3)
add_body(
    '酒店管理模块允许管理员对酒店信息进行增删改查操作。管理员可以添加酒店的基本信息，'
    '管理酒店的房型信息，包括房型名称、价格、床型、设施等。'
)
add_placeholder('【此处插入酒店管理页面截图】')

add_heading_custom('5.3.3 特产管理', level=3)
add_body(
    '特产管理模块包括特产分类管理和商品管理两个部分。管理员可以添加和修改特产分类，'
    '管理特产的名称、价格、库存、描述、图片等信息，以及上架和下架商品。'
)
add_placeholder('【此处插入特产管理页面截图】')

add_heading_custom('5.3.4 用户管理', level=3)
add_body(
    '用户管理模块用于管理系统的注册用户。管理员可以查看用户列表，添加新用户，编辑用户信息，'
    '以及禁用或启用用户账户。该模块有助于维护平台用户秩序。'
)
add_placeholder('【此处插入用户管理页面截图】')

add_heading_custom('5.3.5 订单管理', level=3)
add_body(
    '订单管理模块显示用户在特产商城提交的所有订单。管理员可以查看订单详情，更新订单状态'
    '（如确认发货、标记完成等），处理用户的售后请求。'
)
add_placeholder('【此处插入订单管理页面截图】')

add_heading_custom('5.3.6 攻略管理', level=3)
add_body(
    '攻略管理模块用于审核和管理用户发布的旅游攻略。管理员可以查看攻略列表，'
    '对攻略进行显示或隐藏操作，确保攻略内容符合平台规范。同时支持攻略标签的管理。'
)
add_placeholder('【此处插入攻略管理页面截图】')

add_heading_custom('5.3.7 评论管理', level=3)
add_body(
    '评论管理模块允许管理员查看和管理用户发表的评论。管理员可以对评论进行审核，'
    '隐藏不适当或违规的评论内容，维护平台良好的社区氛围。'
)
add_placeholder('【此处插入评论管理页面截图】')

add_heading_custom('5.3.8 日志管理', level=3)
add_body(
    '日志管理模块记录用户在平台上的重要操作，包括注册、登录、评论、下单、收藏等行为。'
    '管理员可以查看操作日志，了解平台运行情况和用户活动，'
    '为系统维护和安全审计提供依据。'
)
add_placeholder('【此处插入日志管理页面截图】')

# ============================================================
# 第6章 系统测试
# ============================================================
doc.add_page_break()
add_heading_custom('第6章 系统测试', level=1, center=True)

add_heading_custom('6.1 测试环境', level=2)
add_body('系统测试环境配置如下：')

headers_te = ['项目', '配置']
rows_te = [
    ['操作系统', 'Windows 10/11'],
    ['CPU', 'Intel Core i5 / AMD Ryzen 5及以上'],
    ['内存', '8GB及以上'],
    ['Web服务器', 'Apache Tomcat 10+'],
    ['数据库', 'MySQL 5.7+'],
    ['浏览器', 'Chrome / Edge / Firefox（最新版本）'],
]
doc.add_paragraph()
add_table_from_data(headers_te, rows_te, col_widths=[4, 10])

add_heading_custom('6.2 测试用例', level=2)
add_body(
    '本系统采用黑盒测试方法，对核心功能进行了功能测试。以下列出部分测试用例及测试结果：'
)

headers_tc = ['测试编号', '测试功能', '测试步骤', '预期结果', '测试结果']
rows_tc = [
    ['TC-001', '用户注册', '填写注册信息，点击注册按钮', '注册成功，跳转至登录页', '通过'],
    ['TC-002', '用户登录', '输入正确用户名和密码', '登录成功，跳转至首页', '通过'],
    ['TC-003', '用户登录（失败）', '输入错误密码', '登录失败，提示错误信息', '通过'],
    ['TC-004', '景点浏览', '点击景点卡片', '显示景点详情信息', '通过'],
    ['TC-005', '地图导航', '点击"到这里"按钮', '显示路线规划和价格', '通过'],
    ['TC-006', '特产下单', '选择商品→提交订单', '订单创建成功', '通过'],
    ['TC-007', '发布攻略', '填写攻略内容并发布', '攻略发布成功并显示', '通过'],
    ['TC-008', '发表评论', '输入评论内容并提交', '评论显示在内容下方', '通过'],
    ['TC-009', '收藏功能', '点击收藏按钮', '收藏成功，可在个人中心查看', '通过'],
    ['TC-010', '景点管理（后台）', '添加/编辑景点信息', '数据更新成功', '通过'],
    ['TC-011', '用户管理（后台）', '禁用用户账户', '该用户无法登录', '通过'],
    ['TC-012', '订单管理（后台）', '修改订单状态', '订单状态更新', '通过'],
]
add_table_from_data(headers_tc, rows_tc, col_widths=[2, 3, 4.5, 3.5, 1.5])

add_body(
    '测试结果表明，系统的各项核心功能均已正确实现，能够满足预期的业务需求。'
    '系统在不同浏览器下均能正常运行，界面兼容性良好。'
)

# ============================================================
# 第7章 总结与展望
# ============================================================
doc.add_page_break()
add_heading_custom('第7章 总结与展望', level=1, center=True)

add_heading_custom('7.1 系统总结', level=2)
add_body(
    '本文设计并实现了宁夏文旅综合服务平台，完成了从需求分析、系统设计到系统实现和测试的全过程。'
    '系统采用B/S架构，使用Java Servlet、JSP、MyBatis和MySQL技术栈进行开发，'
    '实现了用户端的前台展示和后台管理功能。'
)
add_body(
    '在功能方面，系统提供了景点浏览与导航、酒店信息查询、特产在线购买、旅游攻略分享、'
    '评论与收藏等核心服务，并集成了高德地图API实现了路线规划功能。'
    '后台管理模块涵盖了景点、酒店、特产、用户、订单、攻略、评论、日志等全方位的管理功能。'
    '系统界面简洁美观，操作流程清晰，能够满足用户和管理员的使用需求。'
)
add_body(
    '在技术方面，系统通过MyBatis框架实现了数据访问的便捷管理，采用SHA-256算法保障用户密码安全，'
    '通过会话管理和角色权限控制确保系统安全。系统的三层架构设计使得各功能模块之间耦合度低，'
    '便于后续的维护和功能扩展。'
)

add_heading_custom('7.2 不足与展望', level=2)
add_body(
    '虽然本系统已实现了基本的功能需求，但仍存在一些不足之处和改进空间：'
)
add_body(
    '（1）支付功能：目前特产商城的订单功能已实现订单流程管理，但尚未接入在线支付接口。'
    '后续可集成支付宝或微信支付，实现在线支付功能，提升用户体验。'
)
add_body(
    '（2）个性化推荐：系统目前缺乏个性化推荐功能。后续可根据用户的浏览历史和收藏偏好，'
    '利用推荐算法为用户推荐感兴趣的景点、酒店或特产。'
)
add_body(
    '（3）移动端适配：虽然Bootstrap框架提供了响应式设计的基础支持，但系统的移动端体验仍有待优化。'
    '后续可开发微信小程序或移动端应用，进一步拓展服务渠道。'
)
add_body(
    '（4）性能优化：随着平台用户量和数据量的增长，系统性能可能会面临挑战。'
    '后续可引入缓存机制（如Redis）、数据库读写分离等技术手段优化系统性能。'
)
add_body(
    '（5）人工智能应用：可以利用人工智能技术实现智能客服、语音导览、图像识别等创新功能，'
    '进一步提升平台的智能化水平和服务质量。'
)

# ============================================================
# 参考文献
# ============================================================
doc.add_page_break()
add_heading_custom('参考文献', level=1, font_size=18, center=True, space_before=24, space_after=12)

refs = [
    '[1] 王志勇. Java Web应用开发技术[M]. 北京: 清华大学出版社, 2020.',
    '[2] 刘鑫. MyBatis从入门到精通[M]. 北京: 电子工业出版社, 2019.',
    '[3] 陈刚. MySQL数据库设计与优化[M]. 北京: 机械工业出版社, 2021.',
    '[4] 张涛. 基于B/S架构的Web应用系统设计与实现[J]. 计算机技术与发展, 2022, 32(5): 120-125.',
    '[5] 李明, 王芳. 智慧旅游平台的研究与设计[J]. 旅游学刊, 2021, 36(8): 45-52.',
    '[6] 赵军. 宁夏旅游资源开发与信息化建设研究[D]. 宁夏大学, 2020.',
    '[7] 孙伟. Bootstrap响应式Web开发实战[M]. 北京: 人民邮电出版社, 2022.',
    '[8] 周杰. JSP Servlet技术详解(第3版)[M]. 北京: 清华大学出版社, 2021.',
    '[9] 高德开放平台. 高德地图JavaScript API开发指南[EB/OL]. https://lbs.amap.com/, 2024.',
    '[10] 吴强. 基于Java的旅游信息服务平台设计与实现[J]. 计算机工程与设计, 2023, 44(3): 789-795.',
    '[11] 陈晓明. 互联网+背景下的智慧旅游发展模式研究[J]. 经济研究导刊, 2022(15): 98-101.',
    '[12] Apache Software Foundation. Tomcat Documentation[EB/OL]. https://tomcat.apache.org/, 2024.',
    '[13] MyBatis Team. MyBatis 3 User Guide[EB/OL]. https://mybatis.org/mybatis-3/, 2024.',
    '[14] 黄丽华. 数据库系统原理与应用教程[M]. 北京: 高等教育出版社, 2020.',
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(-0.74)
    pf.left_indent = Cm(0.74)

# ============================================================
# 致谢
# ============================================================
doc.add_page_break()
add_heading_custom('致  谢', level=1, font_size=18, center=True, space_before=24, space_after=12)

add_body(
    '时光荏苒，岁月如梭。转眼间，大学生活即将画上句号。在此毕业论文完成之际，'
    '我谨向所有关心、帮助和支持我的人致以最诚挚的谢意。'
)
add_body(
    '首先，衷心感谢我的指导老师。在论文选题、方案设计、系统开发到论文撰写的整个过程中，'
    '老师都给予了悉心指导和宝贵建议。老师严谨的治学态度、丰富的专业知识和认真负责的工作作风，'
    '深深地感染和激励着我，使我受益匪浅。'
)
add_body(
    '其次，感谢我的同学们。在项目开发和论文撰写过程中，大家互相交流、共同探讨，'
    '在遇到困难时给予了我无私的帮助和鼓励。这段共同奋斗的经历是我人生中宝贵的财富。'
)
add_body(
    '同时，感谢我的家人。感谢他们多年来的默默付出和一如既往的支持与理解，'
    '正是他们的鼓励让我能够全身心地投入到学习和研究中。'
)
add_body(
    '最后，感谢所有参考文献的作者们，他们的研究成果为本文提供了重要的理论依据和参考。'
    '感谢在百忙之中参与论文评审和答辩的各位老师，感谢您们提出的宝贵意见和建议。'
)
add_body(
    '由于本人学识有限，论文中难免存在不足之处，恳请各位老师和专家批评指正。'
)

# ============================================================
# 保存文档
# ============================================================
output_path = 'C:\\Users\\Lenovo\\Desktop\\shixun\\论文_文旅综合服务平台.docx'
doc.save(output_path)
print(f'论文文档已生成: {output_path}')
print(f'文件大小: {os.path.getsize(output_path) / 1024:.1f} KB')
